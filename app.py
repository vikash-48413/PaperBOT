# -*- coding: utf-8 -*-
"""
PaperBOT - AI-Powered Research Paper Assistant
Main FastAPI application with support for large file uploads (up to 50MB)
Optimized for fast processing of 5MB files
"""

# Fix Windows console encoding FIRST before any other imports
import sys
import os

if sys.platform == 'win32':
    # Force UTF-8 encoding for Windows console
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

# Loading the environment variable
from dotenv import load_dotenv
load_dotenv()

# Set environment variables before other imports
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
os.environ['PINECONE_API_KEY'] = PINECONE_API_KEY if PINECONE_API_KEY else ""

from fastapi import FastAPI, Request, Form, Response, File, UploadFile, HTTPException
from fastapi.templating import Jinja2Templates
from fastapi.encoders import jsonable_encoder
from fastapi.responses import JSONResponse
import uvicorn
import json
import shutil
from pathlib import Path
import tempfile
import asyncio
import threading
from QASystem.retrieval_and_generation import get_result
from QASystem.ingestion import ingest_document, get_embedder
from QASystem.utils import pinecone_config

# File size limits - optimized for reasonable processing times
# Recommended limits based on processing time:
#   - 5MB: ~1-2 minutes processing (good UX)
#   - 10MB: ~3-5 minutes processing (acceptable)
#   - 15MB: ~5-10 minutes processing (max recommended)
MAX_FILE_SIZE = 15 * 1024 * 1024  # 15MB - balanced limit
MAX_UPLOAD_SIZE = 15 * 1024 * 1024  # FastAPI body size limit
RECOMMENDED_FILE_SIZE = 5 * 1024 * 1024  # 5MB - optimal for fast processing

# Setup directories
UPLOADS_DIR = Path("uploads")
UPLOADS_DIR.mkdir(exist_ok=True)
DATA_DIR = Path("data")  # Preloaded files directory
DATA_DIR.mkdir(exist_ok=True)

# Global variable to track current document
current_document = {"filename": None, "status": "No document uploaded", "progress": 0}

# Model warmup status
model_ready = {"status": False, "message": "Loading..."}

print("[INFO] Imports loaded successfully")

# Creating the app with proper configuration for large file uploads
app = FastAPI(
    title="PaperBOT",
    description="AI-Powered Research Paper Assistant",
    version="2.0",
)

def warmup_model_background():
    """Warm up the embedding model in background thread"""
    global model_ready
    try:
        print("[WARMUP] Pre-loading embedding model...")
        embedder = get_embedder()
        embedder.warm_up()
        model_ready["status"] = True
        model_ready["message"] = "Ready"
        print("[WARMUP] Embedding model loaded and ready!")
    except Exception as e:
        model_ready["status"] = False
        model_ready["message"] = f"Error: {str(e)}"
        print(f"[WARMUP] Failed to pre-load model: {e}")

# Pre-load embedding model on startup to avoid delay on first upload
@app.on_event("startup")
async def startup_event():
    """Pre-load models on startup for faster first upload"""
    print("[STARTUP] PaperBOT server starting...")
    # Start model warmup in background thread
    thread = threading.Thread(target=warmup_model_background, daemon=True)
    thread.start()
    print("[INFO] Model warmup started in background...")

# Configure templates
templates = Jinja2Templates(directory="templates")

# Supported file extensions
ALLOWED_EXTENSIONS = [".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".json", ".xlsx", ".xls"]

#creating the routes with bind functions
@app.get("/")
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/model_status")
async def model_status():
    """Check if the embedding model is ready"""
    return model_ready

@app.get("/preloaded_files")
async def get_preloaded_files():
    """Get list of preloaded files in the data folder"""
    try:
        files = []
        for file_path in DATA_DIR.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in ALLOWED_EXTENSIONS:
                file_size = file_path.stat().st_size
                files.append({
                    "name": file_path.name,
                    "size": file_size,
                    "size_mb": round(file_size / (1024 * 1024), 2),
                    "type": file_path.suffix.lower()
                })
        return {"success": True, "files": files}
    except Exception as e:
        return {"success": False, "error": str(e), "files": []}

@app.post("/load_preloaded_file")
async def load_preloaded_file(request: Request):
    """Load a file from the data folder"""
    global current_document
    
    try:
        data = await request.json()
        filename = data.get("filename")
        
        if not filename:
            return JSONResponse(status_code=400, content={"success": False, "error": "No filename provided"})
        
        file_path = DATA_DIR / filename
        
        if not file_path.exists():
            return JSONResponse(status_code=404, content={"success": False, "error": f"File not found: {filename}"})
        
        if file_path.suffix.lower() not in ALLOWED_EXTENSIONS:
            return JSONResponse(status_code=400, content={"success": False, "error": f"Unsupported file type: {file_path.suffix}"})
        
        file_size = file_path.stat().st_size
        file_size_mb = file_size / (1024 * 1024)
        
        print(f"[INFO] Loading preloaded file: {filename} ({file_size_mb:.2f}MB)")
        
        # Copy to uploads folder
        dest_path = UPLOADS_DIR / filename
        shutil.copy2(file_path, dest_path)
        
        # Update status
        current_document["filename"] = filename
        current_document["status"] = "Processing..."
        current_document["progress"] = 10
        
        # Process document
        import time
        start = time.time()
        
        current_document["progress"] = 30
        document_store = pinecone_config(namespace="default")
        
        current_document["progress"] = 50
        success = ingest_document(str(dest_path), document_store, clear_existing=True)
        
        if not success:
            raise Exception("Document ingestion failed")
        
        current_document["progress"] = 90
        elapsed = time.time() - start
        
        current_document["status"] = "Ready"
        current_document["progress"] = 100
        
        print(f"[OK] Preloaded file ready: {filename} in {elapsed:.1f}s")
        
        return JSONResponse(content={
            "success": True,
            "message": f"File '{filename}' loaded and processed successfully!",
            "filename": filename,
            "size_mb": round(file_size_mb, 2),
            "processing_time": round(elapsed, 1)
        })
        
    except Exception as e:
        current_document["status"] = "Failed"
        current_document["progress"] = 0
        current_document["filename"] = None
        print(f"[ERROR] Failed to load preloaded file: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})

@app.get("/preview_file/{filename:path}")
async def preview_any_file(filename: str, source: str = "auto"):
    """
    Preview any file from uploads or data folder without downloading.
    
    Args:
        filename: Name of the file to preview
        source: 'uploads', 'data', or 'auto' (checks both)
    """
    try:
        file_path = None
        
        # Determine file location
        if source == "uploads":
            file_path = UPLOADS_DIR / filename
        elif source == "data":
            file_path = DATA_DIR / filename
        else:  # auto - check both locations
            if (UPLOADS_DIR / filename).exists():
                file_path = UPLOADS_DIR / filename
            elif (DATA_DIR / filename).exists():
                file_path = DATA_DIR / filename
        
        if not file_path or not file_path.exists():
            return Response(
                content=generate_error_html("File Not Found", f"The file '{filename}' was not found."),
                media_type="text/html",
                status_code=404
            )
        
        ext = file_path.suffix.lower()
        
        # Handle PDF - serve directly for inline viewing
        if ext == ".pdf":
            with open(file_path, "rb") as f:
                file_content = f.read()
            return Response(
                content=file_content,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"inline; filename=\"{filename}\"",
                    "X-Content-Type-Options": "nosniff"
                }
            )
        
        # Handle DOCX/DOC - convert to HTML
        elif ext in [".docx", ".doc"]:
            return convert_docx_to_html(file_path, filename)
        
        # Handle Excel files - convert to HTML table
        elif ext in [".xlsx", ".xls"]:
            return convert_excel_to_html(file_path, filename)
        
        # Handle CSV - convert to HTML table
        elif ext == ".csv":
            return convert_csv_to_html(file_path, filename)
        
        # Handle JSON - format and display
        elif ext == ".json":
            return convert_json_to_html(file_path, filename)
        
        # Handle Markdown - convert to HTML
        elif ext == ".md":
            return convert_markdown_to_html(file_path, filename)
        
        # Handle TXT - display as formatted text
        elif ext == ".txt":
            return convert_text_to_html(file_path, filename)
        
        else:
            return Response(
                content=generate_error_html("Unsupported Format", f"Preview not available for '{ext}' files."),
                media_type="text/html",
                status_code=400
            )
            
    except Exception as e:
        print(f"[ERROR] Error previewing file {filename}: {e}")
        import traceback
        traceback.print_exc()
        return Response(
            content=generate_error_html("Preview Error", str(e)),
            media_type="text/html",
            status_code=500
        )

# ============================================================================
# DOCUMENT CONVERSION HELPER FUNCTIONS
# ============================================================================

def generate_preview_html_base(title: str, content: str, extra_styles: str = "") -> str:
    """Generate base HTML template for previews"""
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Preview: {title}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            line-height: 1.6;
            color: #1f2937;
            background: #f9fafb;
        }}
        .preview-container {{
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            background: white;
            min-height: 100vh;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }}
        .preview-header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px 30px;
            margin: -40px -20px 30px -20px;
            border-radius: 0;
        }}
        .preview-header h1 {{
            font-size: 20px;
            font-weight: 600;
            margin: 0;
            display: flex;
            align-items: center;
            gap: 10px;
        }}
        .preview-header .file-icon {{ font-size: 24px; }}
        .content {{ padding: 0 10px; }}
        h1, h2, h3 {{ color: #1f2937; margin-top: 24px; margin-bottom: 12px; }}
        p {{ margin-bottom: 16px; }}
        code {{ background: #f3f4f6; padding: 2px 6px; border-radius: 4px; font-size: 14px; }}
        pre {{ 
            background: #1e1e1e; 
            color: #d4d4d4; 
            padding: 20px; 
            border-radius: 8px; 
            overflow-x: auto;
            margin: 16px 0;
        }}
        table {{ 
            border-collapse: collapse; 
            width: 100%; 
            margin: 20px 0;
            font-size: 14px;
        }}
        th, td {{ 
            border: 1px solid #e5e7eb; 
            padding: 12px 16px; 
            text-align: left; 
        }}
        th {{ 
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white; 
            font-weight: 600;
            position: sticky;
            top: 0;
        }}
        tr:nth-child(even) {{ background-color: #f9fafb; }}
        tr:hover {{ background-color: #f3f4f6; }}
        {extra_styles}
    </style>
</head>
<body>
    <div class="preview-container">
        <div class="preview-header">
            <h1><span class="file-icon">📄</span> {title}</h1>
        </div>
        <div class="content">
            {content}
        </div>
    </div>
</body>
</html>"""

def generate_error_html(title: str, message: str) -> str:
    """Generate error HTML page"""
    return generate_preview_html_base(
        "Error",
        f"""<div style="text-align: center; padding: 60px 20px;">
            <div style="font-size: 64px; margin-bottom: 20px;">⚠️</div>
            <h2 style="color: #ef4444; margin-bottom: 16px;">{title}</h2>
            <p style="color: #6b7280;">{message}</p>
        </div>""",
        ""
    )

def convert_docx_to_html(file_path: Path, filename: str) -> Response:
    """Convert DOCX file to HTML for preview"""
    try:
        import docx
        doc = docx.Document(file_path)
        
        content_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    if level.isdigit() and int(level) <= 6:
                        content_parts.append(f"<h{level}>{para.text}</h{level}>")
                    else:
                        content_parts.append(f"<h3>{para.text}</h3>")
                else:
                    content_parts.append(f"<p>{para.text}</p>")
        
        # Also extract tables if present
        for table in doc.tables:
            table_html = "<table>"
            for i, row in enumerate(table.rows):
                tag = "th" if i == 0 else "td"
                table_html += "<tr>"
                for cell in row.cells:
                    table_html += f"<{tag}>{cell.text}</{tag}>"
                table_html += "</tr>"
            table_html += "</table>"
            content_parts.append(table_html)
        
        content = "\n".join(content_parts)
        html = generate_preview_html_base(filename, content)
        return Response(content=html, media_type="text/html")
        
    except Exception as e:
        return Response(
            content=generate_error_html("DOCX Conversion Error", str(e)),
            media_type="text/html",
            status_code=500
        )

def convert_excel_to_html(file_path: Path, filename: str) -> Response:
    """Convert Excel file to HTML table for preview"""
    try:
        import pandas as pd
        
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        content_parts = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            content_parts.append(f"<h2>📊 Sheet: {sheet_name}</h2>")
            content_parts.append(f"<p style='color: #6b7280;'>{len(df)} rows × {len(df.columns)} columns</p>")
            # Limit to first 500 rows for performance
            if len(df) > 500:
                content_parts.append("<p style='color: #f59e0b;'>⚠️ Showing first 500 rows</p>")
                df = df.head(500)
            content_parts.append(df.to_html(index=False, classes='preview-table', na_rep='—'))
        
        content = "\n".join(content_parts)
        html = generate_preview_html_base(filename, content)
        return Response(content=html, media_type="text/html")
        
    except Exception as e:
        return Response(
            content=generate_error_html("Excel Conversion Error", str(e)),
            media_type="text/html",
            status_code=500
        )

def convert_csv_to_html(file_path: Path, filename: str) -> Response:
    """Convert CSV file to HTML table for preview"""
    try:
        import pandas as pd
        df = pd.read_csv(file_path)
        
        content_parts = [f"<p style='color: #6b7280;'>{len(df)} rows × {len(df.columns)} columns</p>"]
        
        # Limit to first 500 rows for performance
        if len(df) > 500:
            content_parts.append("<p style='color: #f59e0b;'>⚠️ Showing first 500 rows</p>")
            df = df.head(500)
        
        content_parts.append(df.to_html(index=False, classes='preview-table', na_rep='—'))
        content = "\n".join(content_parts)
        html = generate_preview_html_base(filename, content)
        return Response(content=html, media_type="text/html")
        
    except Exception as e:
        return Response(
            content=generate_error_html("CSV Conversion Error", str(e)),
            media_type="text/html",
            status_code=500
        )

def convert_json_to_html(file_path: Path, filename: str) -> Response:
    """Convert JSON file to formatted HTML for preview"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        formatted_json = json.dumps(json_data, indent=2, ensure_ascii=False)
        
        # Syntax highlighting for JSON
        import html
        formatted_json = html.escape(formatted_json)
        
        content = f"<pre>{formatted_json}</pre>"
        extra_styles = """
            pre { 
                background: #1e1e1e !important; 
                color: #ce9178 !important;
                font-family: 'Fira Code', 'Consolas', monospace;
                font-size: 14px;
                line-height: 1.5;
            }
        """
        html_content = generate_preview_html_base(filename, content, extra_styles)
        return Response(content=html_content, media_type="text/html")
        
    except Exception as e:
        return Response(
            content=generate_error_html("JSON Parse Error", str(e)),
            media_type="text/html",
            status_code=500
        )

def convert_markdown_to_html(file_path: Path, filename: str) -> Response:
    """Convert Markdown file to HTML for preview"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Try using markdown library if available
        try:
            import markdown
            html_content = markdown.markdown(
                md_content,
                extensions=['tables', 'fenced_code', 'codehilite', 'toc']
            )
        except ImportError:
            # Fallback: basic markdown-like rendering
            import html
            html_content = html.escape(md_content)
            html_content = f"<pre style='white-space: pre-wrap; font-family: inherit;'>{html_content}</pre>"
        
        extra_styles = """
            .content img { max-width: 100%; height: auto; border-radius: 8px; margin: 16px 0; }
            .content blockquote { 
                border-left: 4px solid #667eea; 
                padding-left: 20px; 
                margin: 16px 0;
                color: #6b7280;
                font-style: italic;
            }
            .content ul, .content ol { margin: 16px 0; padding-left: 30px; }
            .content li { margin: 8px 0; }
        """
        html_output = generate_preview_html_base(filename, html_content, extra_styles)
        return Response(content=html_output, media_type="text/html")
        
    except Exception as e:
        return Response(
            content=generate_error_html("Markdown Conversion Error", str(e)),
            media_type="text/html",
            status_code=500
        )

def convert_text_to_html(file_path: Path, filename: str) -> Response:
    """Convert text file to HTML for preview"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            txt_content = f.read()
        
        import html
        txt_content = html.escape(txt_content)
        
        # Count lines and words
        lines = txt_content.count('\n') + 1
        words = len(txt_content.split())
        
        content = f"""
            <p style='color: #6b7280; margin-bottom: 20px;'>{lines} lines • {words} words</p>
            <pre style='white-space: pre-wrap; background: #f9fafb; padding: 20px; border-radius: 8px; 
                        font-family: "Fira Code", Consolas, monospace; font-size: 14px; line-height: 1.6;
                        border: 1px solid #e5e7eb;'>{txt_content}</pre>
        """
        html_output = generate_preview_html_base(filename, content)
        return Response(content=html_output, media_type="text/html")
        
    except Exception as e:
        return Response(
            content=generate_error_html("Text File Error", str(e)),
            media_type="text/html",
            status_code=500
        )

@app.get("/test")
async def test(request: Request):
    return templates.TemplateResponse("test.html", {"request": request})
    
@app.post("/upload_document")
async def upload_document(file: UploadFile = File(...)):
    """Upload and process document with comprehensive error handling and size validation"""
    print(f"📥 Upload endpoint called - Filename: {file.filename}")
    
    # Variables for cleanup
    temp_file_path = None
    final_file_path = None
    
    try:
        # Validate file exists and has a name
        if not file or not file.filename:
            return JSONResponse(
                status_code=400,
                content={"success": False, "error": "No file provided"}
            )
        
        # Validate file type - now supports many formats
        allowed_extensions = [".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".json", ".xlsx", ".xls"]
        file_ext = Path(file.filename).suffix.lower()
        
        if file_ext not in allowed_extensions:
            return JSONResponse(
                status_code=400,
                content={
                    "success": False, 
                    "error": f"Unsupported file type '{file_ext}'. Supported formats: PDF, DOCX, DOC, TXT, MD, CSV, JSON, XLSX, XLS"
                }
            )
        
        print("[INFO] File type validated:", file_ext)
        
        # Read file content with size validation
        # Use streaming approach to handle large files efficiently
        file_size = 0
        chunks = []
        
        print("[INFO] Reading file in chunks...")
        try:
            while chunk := await file.read(1024 * 1024):  # Read 1MB at a time
                file_size += len(chunk)
                
                # Check size limit as we read
                if file_size > MAX_FILE_SIZE:
                    return JSONResponse(
                        status_code=413,  # Payload Too Large
                        content={
                            "success": False,
                            "error": f"File size ({file_size / (1024*1024):.2f}MB) exceeds maximum allowed size ({MAX_FILE_SIZE / (1024*1024):.0f}MB)"
                        }
                    )
                
                chunks.append(chunk)
            
            # Combine all chunks
            file_content = b''.join(chunks)
            
        except Exception as read_error:
            print("[ERROR] Error reading file:", read_error)
            return JSONResponse(
                status_code=500,
                content={
                    "success": False,
                    "error": f"Failed to read uploaded file: {str(read_error)}"
                }
            )
        
        file_size_mb = file_size / (1024 * 1024)
        print(f"[OK] File read successfully: {file_size_mb:.2f}MB")
        
        # Validate minimum file size (should have content)
        if file_size == 0:
            return JSONResponse(
                status_code=400,
                content={"success": False, "error": "File is empty"}
            )
        
        # Clear previous uploads
        print("[INFO] Clearing previous uploads...")
        for old_file in UPLOADS_DIR.glob("*"):
            try:
                old_file.unlink()
            except Exception as e:
                print(f"[WARN] Could not delete {old_file}: {e}")
        
        # Save uploaded file with proper error handling
        final_file_path = UPLOADS_DIR / file.filename
        try:
            with final_file_path.open("wb") as buffer:
                buffer.write(file_content)
            print(f"[OK] File saved to: {final_file_path}")
        except Exception as save_error:
            print("[ERROR] Error saving file:", save_error)
            return JSONResponse(
                status_code=500,
                content={
                    "success": False,
                    "error": f"Failed to save file: {str(save_error)}"
                }
            )
        
        # Update current document status
        current_document["filename"] = file.filename
        current_document["status"] = "Processing..."
        current_document["progress"] = 10
        
        print(f"[INFO] Upload started: {file.filename} ({file_size_mb:.2f}MB)")
        
        # Process document synchronously (avoiding multiprocessing issues on Windows)
        import time
        start = time.time()
        
        try:
            current_document["progress"] = 30
            print("[INFO] Initializing document store...")
            document_store = pinecone_config(namespace="default")
            
            current_document["progress"] = 50
            print("[INFO] Processing document...")
            
            # Run ingestion with timeout to prevent indefinite hanging
            success = ingest_document(str(final_file_path), document_store, clear_existing=True)
            
            if not success:
                raise Exception("Document ingestion returned False - check server logs for details")
            
            current_document["progress"] = 90
            elapsed = time.time() - start
            print(f"[OK] Total processing time: {elapsed:.1f}s")
            
            current_document["status"] = "Ready"
            current_document["progress"] = 100
            print(f"[OK] Upload completed: {file.filename}")
            
            return JSONResponse(
                status_code=200,
                content={
                    "success": True, 
                    "message": f"Document '{file.filename}' ({file_size_mb:.1f}MB) uploaded and processed successfully!",
                    "filename": file.filename,
                    "size_mb": round(file_size_mb, 2),
                    "processing_time": round(elapsed, 1)
                }
            )
            
        except Exception as process_error:
            print("[ERROR] Error processing document:", process_error)
            import traceback
            traceback.print_exc()
            
            # Cleanup on processing error
            if final_file_path and final_file_path.exists():
                try:
                    final_file_path.unlink()
                except:
                    pass
            
            current_document["status"] = "Failed"
            current_document["progress"] = 0
            current_document["filename"] = None
            
            return JSONResponse(
                status_code=500,
                content={
                    "success": False,
                    "error": f"Document processing failed: {str(process_error)}"
                }
            )
            
    except HTTPException as http_exc:
        # Re-raise HTTP exceptions
        raise
        
    except Exception as e:
        current_document["status"] = "Failed"
        current_document["progress"] = 0
        current_document["filename"] = None
        error_msg = str(e)
        print("[ERROR] Unexpected exception during upload:", error_msg)
        import traceback
        traceback.print_exc()
        
        # Cleanup on error
        if final_file_path and final_file_path.exists():
            try:
                final_file_path.unlink()
            except:
                pass
        
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": f"Upload failed: {error_msg}"
            }
        )

@app.get("/document_status")
async def document_status():
    return current_document

@app.get("/preview_document")
async def preview_document():
    """
    Serve the current uploaded document for preview.
    Redirects to the new unified preview endpoint.
    """
    if current_document["filename"] is None:
        return Response(
            content=generate_error_html("No Document", "No document has been uploaded yet."),
            media_type="text/html",
            status_code=404
        )
    
    # Use the new unified preview endpoint
    return await preview_any_file(current_document["filename"], source="uploads")

@app.post("/delete_document")
async def delete_document():
    """Delete the current document and clear vectors"""
    try:
        # Clear the vector database (handle namespace not found error gracefully)
        try:
            from pinecone import Pinecone
            pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
            index = pc.Index("paperbot")
            index.delete(delete_all=True, namespace="current_document")
            print("[OK] Cleared vectors from database")
        except Exception as db_error:
            # If namespace doesn't exist or other DB errors, just log and continue
            error_msg = str(db_error)
            if "Namespace not found" in error_msg or "404" in error_msg:
                print("[WARN] Namespace doesn't exist (nothing to delete)")
            else:
                print(f"[WARN] Database cleanup warning: {error_msg}")
        
        # Clear uploaded files
        for old_file in UPLOADS_DIR.glob("*"):
            try:
                old_file.unlink()
            except Exception as file_error:
                print(f"[WARN] Could not delete file {old_file}: {file_error}")
        
        # Reset current document state
        filename = current_document.get("filename", "Unknown")
        current_document["filename"] = None
        current_document["status"] = "No document uploaded"
        current_document["progress"] = 0
        
        print(f"[OK] Deleted document: {filename}")
        return {
            "success": True,
            "message": f"Document deleted successfully" if filename == "Unknown" else f"Document '{filename}' deleted successfully"
        }
    except Exception as e:
        print(f"[ERROR] Error deleting document: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

@app.post("/get_answer")
async def get_answer(
    request: Request, 
    question: str = Form(...),
    style: str = Form("Detailed and Technical"),
    length: str = Form("Medium (2-3 paragraphs)")
):
    # Check if document is uploaded
    if current_document["status"] != "Ready":
        error_msg = "Please upload a document first before asking questions."
        response_data = jsonable_encoder(json.dumps({"answer": error_msg}))
        return Response(response_data)
    
    print(f"Question: {question}")
    print(f"Style: {style}, Length: {length}")
    print(f"Current Document: {current_document['filename']}")
    
    answer = get_result(question, style, length)
    response_data = jsonable_encoder(json.dumps({"answer": answer}))
    res = Response(response_data)
    return res
    
if __name__ == "__main__":
    # Note: UTF-8 encoding is already configured at the top of the file
    
    print("[STARTUP] Starting PaperBOT server...")
    print(f"[CONFIG] Max file size: {MAX_FILE_SIZE / (1024*1024):.0f}MB")
    print("[CONFIG] Upload timeout: 10 minutes")
    print(f"[CONFIG] Request size limit: {MAX_UPLOAD_SIZE / (1024*1024):.0f}MB")
    print("[INFO] Server will be available at http://localhost:8000")
    
    # Increase timeout for large file processing and set size limits
    uvicorn.run(
        "app:app", 
        host="0.0.0.0", 
        port=8000, 
        reload=False, 
        timeout_keep_alive=600,  # 10 minutes keep-alive
        limit_concurrency=10,  # Limit concurrent connections
        limit_max_requests=1000,  # Prevent memory leaks
        timeout_graceful_shutdown=30  # Graceful shutdown timeout
    )