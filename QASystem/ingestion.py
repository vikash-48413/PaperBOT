from haystack import Pipeline, Document
from haystack.components.writers import DocumentWriter
from haystack.components.preprocessors import DocumentSplitter
from haystack.components.embedders import SentenceTransformersDocumentEmbedder
from haystack_integrations.document_stores.pinecone import PineconeDocumentStore
from haystack.components.converters import PyPDFToDocument, TextFileToDocument
from pathlib import Path
import os
import gc
import psutil
import time
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
from typing import List, Optional
from dotenv import load_dotenv
from QASystem.utils import pinecone_config
from pinecone import Pinecone

try:
    from QASystem.config import (CURRENT_MODEL, EMBEDDING_MODELS, CHUNK_SETTINGS, BATCH_SIZE,
                                  LARGE_FILE_THRESHOLD, LARGE_FILE_CHUNK_LENGTH, LARGE_FILE_BATCH_SIZE,
                                  MAX_CHUNK_CHARS)
except ImportError:
    # Fallback if config not available
    CURRENT_MODEL = "quality"
    EMBEDDING_MODELS = {"quality": {"name": "BAAI/bge-large-en-v1.5", "dimension": 1024, "description": "Best quality"}}
    CHUNK_SETTINGS = {"split_by": "word", "split_length": 200, "split_overlap": 20}
    BATCH_SIZE = 16
    LARGE_FILE_THRESHOLD = 3 * 1024 * 1024
    LARGE_FILE_CHUNK_LENGTH = 250
    LARGE_FILE_BATCH_SIZE = 20
    MAX_CHUNK_CHARS = 8000  # Conservative limit for Pinecone

# Pinecone metadata size limit - use 8KB to be safe (actual limit is 40KB but includes all metadata)
PINECONE_METADATA_LIMIT = 8000

# Cache the embedder model to avoid reloading (significant speedup)
_cached_embedder = None

def get_memory_usage():
    """Get current memory usage in MB"""
    process = psutil.Process()
    return process.memory_info().rss / 1024 / 1024

def clear_memory():
    """Force garbage collection to free memory"""
    gc.collect()

def enforce_chunk_size_limit(chunks: List[Document], max_chars: int = None) -> List[Document]:
    """
    Ensure all chunks are under Pinecone's metadata size limit.
    Splits oversized chunks into smaller pieces.
    """
    if max_chars is None:
        max_chars = PINECONE_METADATA_LIMIT
    
    valid_chunks = []
    for chunk in chunks:
        content = chunk.content
        if len(content) <= max_chars:
            valid_chunks.append(chunk)
        else:
            # Split oversized chunk into smaller pieces
            print(f"   [WARN] Chunk too large ({len(content)} chars), splitting...")
            for i in range(0, len(content), max_chars):
                sub_content = content[i:i + max_chars]
                if sub_content.strip():  # Only add non-empty chunks
                    new_chunk = Document(
                        content=sub_content,
                        meta={**chunk.meta, "sub_chunk": i // max_chars}
                    )
                    valid_chunks.append(new_chunk)
    
    return valid_chunks

def get_embedder():
    """Get or create cached embedder instance for performance"""
    global _cached_embedder
    if _cached_embedder is None:
        model_config = EMBEDDING_MODELS[CURRENT_MODEL]
        print(f"[INFO] Loading {CURRENT_MODEL} embedding model: {model_config['name']}")
        print(f"   {model_config['description']}")
        print(f"   Dimension: {model_config['dimension']}")
        try:
            _cached_embedder = SentenceTransformersDocumentEmbedder(
                model=model_config['name'],
                batch_size=BATCH_SIZE,
                progress_bar=False,  # Disable progress bar to reduce overhead
                normalize_embeddings=True
                # device is auto-detected (CPU or GPU based on availability)
            )
            print("   [OK] Model loaded successfully")
        except Exception as e:
            print(f"   [ERROR] Failed to load embedding model: {e}")
            raise
    return _cached_embedder

def convert_docx_to_text(file_path: str) -> List[Document]:
    """Convert DOCX to Haystack Document"""
    try:
        import docx
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        text_content = '\n\n'.join(full_text)
        return [Document(content=text_content, meta={"file_path": file_path, "file_type": "docx"})]
    except Exception as e:
        print(f"Error converting DOCX: {e}")
        return []

def convert_csv_to_text(file_path: str) -> List[Document]:
    """Convert CSV to Haystack Document"""
    try:
        import pandas as pd
        df = pd.read_csv(file_path)
        text_content = df.to_string(index=False)
        return [Document(content=text_content, meta={"file_path": file_path, "file_type": "csv"})]
    except Exception as e:
        print(f"Error converting CSV: {e}")
        return []

def convert_excel_to_text(file_path: str) -> List[Document]:
    """Convert Excel to Haystack Document"""
    try:
        import pandas as pd
        df = pd.read_excel(file_path)
        text_content = df.to_string(index=False)
        return [Document(content=text_content, meta={"file_path": file_path, "file_type": "excel"})]
    except Exception as e:
        print(f"Error converting Excel: {e}")
        return []

def convert_json_to_text(file_path: str) -> List[Document]:
    """Convert JSON to Haystack Document"""
    try:
        import json
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        text_content = json.dumps(data, indent=2)
        return [Document(content=text_content, meta={"file_path": file_path, "file_type": "json"})]
    except Exception as e:
        print(f"Error converting JSON: {e}")
        return []

def convert_markdown_to_text(file_path: str) -> List[Document]:
    """Convert Markdown to Haystack Document"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        return [Document(content=text_content, meta={"file_path": file_path, "file_type": "markdown"})]
    except Exception as e:
        print(f"Error converting Markdown: {e}")
        return []

def get_converter_for_file(file_path):
    """Return appropriate converter based on file extension"""
    ext = Path(file_path).suffix.lower()
    
    if ext == '.pdf':
        return PyPDFToDocument()
    elif ext in ['.txt']:
        return TextFileToDocument()
    elif ext == '.md':
        return None  # Handle with custom converter
    elif ext in ['.docx', '.doc']:
        return None  # Handle with custom converter
    elif ext == '.csv':
        return None  # Handle with custom converter
    elif ext in ['.xlsx', '.xls']:
        return None  # Handle with custom converter
    elif ext == '.json':
        return None  # Handle with custom converter
    else:
        return None

def process_chunks_parallel(chunks: List[Document], embedder, batch_size: int = 32) -> List[Document]:
    """Process chunks in parallel batches for faster embedding with comprehensive error handling"""
    embedded_chunks = []
    total_batches = (len(chunks) + batch_size - 1) // batch_size
    
    if not chunks:
        print("   [WARN] No chunks to process")
        return []
    
    print(f"   [INFO] Processing {len(chunks)} chunks in {total_batches} batches (batch_size={batch_size})")
    
    failed_batches = []
    
    try:
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            batch_num = (i // batch_size) + 1
            
            # Monitor memory
            mem_before = get_memory_usage()
            
            try:
                # Validate batch has content
                if not batch:
                    print(f"   [WARN] Skipping empty batch {batch_num}")
                    continue
                
                # Embed batch
                result = embedder.run(documents=batch)
                
                if not result or 'documents' not in result:
                    raise Exception(f"Embedder returned invalid result for batch {batch_num}")
                
                embedded_batch = result['documents']
                
                # Validate embeddings were created
                if not embedded_batch:
                    raise Exception(f"No embeddings generated for batch {batch_num}")
                
                embedded_chunks.extend(embedded_batch)
                
                mem_after = get_memory_usage()
                mem_delta = mem_after - mem_before
                
                print(f"   [OK] Batch {batch_num}/{total_batches} complete ({len(embedded_batch)} chunks, +{mem_delta:.1f}MB)")
                
                # Clear memory periodically for large files
                if batch_num % 3 == 0:
                    clear_memory()
                    print(f"   [GC] Memory cleanup: {get_memory_usage():.1f}MB")
                    
            except Exception as batch_error:
                error_msg = str(batch_error)
                print(f"   [ERROR] Error in batch {batch_num}: {error_msg}")
                failed_batches.append(batch_num)
                
                # If too many batches fail, stop processing
                if len(failed_batches) > total_batches * 0.2:  # More than 20% failure
                    raise Exception(f"Too many batch failures ({len(failed_batches)}/{batch_num}). Aborting.")
                
                # Continue with next batch for isolated failures
                continue
        
        # Report results
        if failed_batches:
            print(f"   [WARN] {len(failed_batches)} batches failed: {failed_batches}")
        
        if not embedded_chunks:
            raise Exception("No chunks were successfully embedded")
        
        return embedded_chunks
        
    except Exception as e:
        print(f"   [FATAL] Error in parallel processing: {str(e)}")
        raise

def ingest_document(file_path, document_store, clear_existing=True, max_pages=None, progress_callback=None):
    """
    Ingest a document into the vector store with parallel processing and memory management.
    
    Args:
        file_path: Path to the document file
        document_store: Pinecone document store instance
        clear_existing: Whether to clear existing documents first
        max_pages: Maximum pages to process (None = all pages)
        progress_callback: Optional callback function for progress updates
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Validate inputs
        if not file_path:
            print("[ERROR] No file path provided")
            return False
        
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            print(f"[ERROR] File does not exist: {file_path}")
            return False
        
        if not document_store:
            print("[ERROR] No document store provided")
            return False
        
        print(f"\n[START] Document ingestion: {file_path_obj.name}")
        initial_memory = get_memory_usage()
        print(f"   [MEM] Server baseline memory: {initial_memory:.1f}MB (includes Python, libraries, and cached models)")
        start_time = time.time()
        
        # Clear existing documents COMPLETELY to avoid mixing different papers
        if clear_existing:
            try:
                load_dotenv()
                pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
                index = pc.Index("paperbot")
                index.delete(delete_all=True, namespace="default")
                print("   [OK] Cleared existing documents from vector store")
            except Exception as e:
                print(f"   [WARN] Error during cleanup: {e}")
                print("   Continuing with ingestion...")
        
        # Get appropriate converter and handle custom formats
        ext = file_path_obj.suffix.lower()
        converter = get_converter_for_file(file_path)
        documents = []
        
        # Convert document to text
        print(f"   [INFO] Converting {ext} file...")
        try:
            if ext in ['.docx', '.doc']:
                documents = convert_docx_to_text(file_path)
            elif ext == '.csv':
                documents = convert_csv_to_text(file_path)
            elif ext in ['.xlsx', '.xls']:
                documents = convert_excel_to_text(file_path)
            elif ext == '.json':
                documents = convert_json_to_text(file_path)
            elif ext == '.md':
                documents = convert_markdown_to_text(file_path)
            elif converter:
                # Use standard Haystack converter for PDF and TXT
                result = converter.run(sources=[file_path_obj])
                documents = result['documents']
            else:
                print(f"   [ERROR] Unsupported file type: {ext}")
                return False
        except Exception as convert_error:
            print(f"   [ERROR] Error converting document: {convert_error}")
            import traceback
            traceback.print_exc()
            return False
        
        if not documents:
            print("   [ERROR] No content extracted from document")
            return False
        
        # Validate document content
        total_content_length = sum(len(doc.content) for doc in documents if doc.content)
        if total_content_length == 0:
            print("   [ERROR] Document has no text content")
            return False
        
        print(f"   [OK] Extracted content ({len(documents)} document(s), {total_content_length} chars)")
        
        # Check file size and optimize settings for large files
        file_size = os.path.getsize(file_path)
        file_size_mb = file_size / (1024 * 1024)
        print(f"   [INFO] File size: {file_size_mb:.2f}MB")
        
        # Optimize chunk settings for large files
        chunk_length = CHUNK_SETTINGS["split_length"]
        split_by = CHUNK_SETTINGS["split_by"]
        batch_size = BATCH_SIZE
        
        if file_size > LARGE_FILE_THRESHOLD:
            chunk_length = LARGE_FILE_CHUNK_LENGTH
            batch_size = LARGE_FILE_BATCH_SIZE
            print("   [OPTIMIZE] Large file detected - using optimized settings:")
            print(f"      Chunk length: {chunk_length} {split_by}s")
            print(f"      Batch size: {batch_size} chunks")
        
        # Split documents into chunks
        print("   [INFO] Splitting into chunks...")
        try:
            splitter = DocumentSplitter(
                split_by=split_by,
                split_length=chunk_length,
                split_overlap=CHUNK_SETTINGS["split_overlap"]
            )
            split_result = splitter.run(documents=documents)
            chunks = split_result['documents']
            
            if not chunks:
                raise Exception("Splitter returned no chunks")
            
            initial_chunk_count = len(chunks)
            print(f"   [OK] Created {initial_chunk_count} chunks from splitter")
            
            # Enforce Pinecone's metadata size limit (40KB per vector)
            chunks = enforce_chunk_size_limit(chunks, PINECONE_METADATA_LIMIT)
            
            if len(chunks) != initial_chunk_count:
                print(f"   [INFO] After size enforcement: {len(chunks)} chunks (was {initial_chunk_count})")
            
        except Exception as split_error:
            print(f"   [ERROR] Error splitting document: {split_error}")
            import traceback
            traceback.print_exc()
            return False
        
        # Embed chunks with parallel processing and memory management
        print("   [INFO] Embedding chunks with parallel processing...")
        try:
            embedder = get_embedder()
            embedder.warm_up()
            
            embedded_chunks = process_chunks_parallel(chunks, embedder, batch_size=batch_size)
            
            # Validate embedding results
            if not embedded_chunks:
                raise Exception("No chunks were successfully embedded. Check embedding model and Pinecone configuration.")
            
            if len(embedded_chunks) != len(chunks):
                print(f"   [WARN] {len(chunks)} chunks created but only {len(embedded_chunks)} embedded")
                # Continue if we have at least some embeddings
                if len(embedded_chunks) < len(chunks) * 0.5:  # Less than 50% success
                    raise Exception(f"Too few chunks embedded ({len(embedded_chunks)}/{len(chunks)})")
            
            print(f"   [OK] Successfully embedded {len(embedded_chunks)} chunks")
            
        except Exception as embed_error:
            print(f"   [ERROR] Error embedding chunks: {embed_error}")
            import traceback
            traceback.print_exc()
            return False
        
        # Write to document store
        print("   [INFO] Writing to vector database...")
        try:
            writer = DocumentWriter(document_store, policy="UPSERT")
            write_result = writer.run(documents=embedded_chunks)
            chunks_written = write_result['documents_written']
            
            if chunks_written == 0:
                raise Exception("No chunks were written to database")
            
            print(f"   [OK] Wrote {chunks_written} chunks to Pinecone")
            
        except Exception as write_error:
            print(f"   [ERROR] Error writing to Pinecone: {str(write_error)}")
            print("   Check your Pinecone API key, index name, and network connection")
            import traceback
            traceback.print_exc()
            return False
        
        # Final memory cleanup
        clear_memory()
        
        elapsed_time = time.time() - start_time
        final_memory = get_memory_usage()
        memory_delta = final_memory - initial_memory
        
        print(f"\n[OK] Ingestion completed successfully!")
        print(f"   Statistics:")
        print(f"   - Document: {file_path_obj.name}")
        print(f"   - Format: {ext.upper()}")
        print(f"   - Size: {file_size_mb:.2f}MB")
        print(f"   - Chunks created: {chunks_written}")
        print(f"   - Time taken: {elapsed_time:.2f} seconds")
        print(f"   - Speed: {chunks_written/elapsed_time:.1f} chunks/sec")
        print(f"   - Memory used for this document: {memory_delta:+.1f}MB (baseline: {initial_memory:.1f}MB -> final: {final_memory:.1f}MB)")
        
        if progress_callback:
            progress_callback(100)
        
        return True
        
    except Exception as e:
        print(f"\n[CRITICAL] Error during ingestion: {str(e)}")
        print(f"   Error type: {type(e).__name__}")
        import traceback
        print("\nFull error trace:")
        traceback.print_exc()
        
        # Cleanup on error
        try:
            clear_memory()
            print("   [OK] Memory cleaned up after error")
        except:
            pass
        
        return False

def ingest(document_store):
    """Legacy function for backward compatibility"""
    default_file = Path(r"E:\PaperBOT\data\Attention_Is_All_Need.pdf")
    return ingest_document(str(default_file), document_store)

if __name__ == '__main__':
    document_store = pinecone_config()
    ingest(document_store)


